"""File upload and processing functionality."""

import json
import pandas as pd
from werkzeug.utils import secure_filename


def extract_tables_from_docx(path):
    """Extract tables from DOCX files"""
    try:
        from docx import Document
        doc = Document(path)
        table_texts = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            table_texts.append("\n".join(rows))
        return "\n\n".join(table_texts)
    except Exception as e:
        return f"[Error extracting tables: {e}]"


def read_extra_file(file_path):
    """Read and process additional context files"""
    try:
        if file_path.endswith(".txt"):
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif file_path.endswith(".json"):
            with open(file_path, 'r', encoding='utf-8') as f:
                return json.dumps(json.load(f), indent=2)
        elif file_path.endswith(".pdf"):
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(file_path)
                text_parts = []
                for page in doc:
                    text_parts.append(page.get_text())
                doc.close()
                return "\n".join(text_parts)
            except ImportError:
                return "[PDF processing requires PyMuPDF library]"
        elif file_path.endswith(".docx"):
            try:
                from docx import Document
                doc = Document(file_path)
                text_parts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
                table_text = extract_tables_from_docx(file_path)
                return "\n\n".join(text_parts + [table_text])
            except ImportError:
                return "[DOCX processing requires python-docx library]"
        elif file_path.endswith(".xlsx"):
            try:
                df = pd.read_excel(file_path)
                return df.to_csv(index=False)
            except Exception as e:
                return f"[Error reading Excel file: {e}]"
        elif file_path.endswith(".csv"):
            try:
                df = pd.read_csv(file_path)
                return df.to_csv(index=False)
            except Exception as e:
                return f"[Error reading CSV file: {e}]"
        else:
            return f"[Unsupported file format: {file_path}]"
    except Exception as e:
        return f"[Error reading extra file: {e}]"


def validate_csv_file(file):
    """Validate CSV file upload"""
    if not file or file.filename == '':
        return False, 'No CSV file selected'

    if not file.filename.lower().endswith('.csv'):
        return False, 'Main file must be a CSV'

    return True, None


def process_csv_file(file):
    """Process uploaded CSV file and return dataset"""
    try:
        dataset = pd.read_csv(file)
        if dataset.empty:
            return None, 'CSV file is empty'
        return dataset, None
    except Exception as e:
        return None, f'Error reading CSV file: {str(e)}'


def validate_extra_file(file):
    """Validate additional context file"""
    if not file or file.filename == '':
        return True, None  # Extra file is optional

    allowed_extensions = ['.txt', '.json', '.pdf', '.docx', '.xlsx', '.csv']
    file_ext = '.' + file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else ''

    if file_ext not in allowed_extensions:
        return False, f'Unsupported extra file format. Allowed: {", ".join(allowed_extensions)}'

    return True, None


def get_file_info(dataset, filename):
    """Get basic information about the uploaded dataset"""
    return {
        'filename': filename,
        'rows': len(dataset),
        'columns': len(dataset.columns),
        'column_names': list(dataset.columns),
        'preview': dataset.head().fillna('').to_dict('records')
    }


def secure_filename_with_session(filename, session_id):
    """Create a secure filename with session prefix"""
    secure_name = secure_filename(filename)
    return f"{session_id}_{secure_name}"