from flask import Flask, render_template, request, jsonify, send_file, Response
import pandas as pd
import json
import os
import tempfile
import numpy as np
from werkzeug.utils import secure_filename
from datetime import datetime

# Import functions from the updated meta_data_ex_api.py
from meta_data_ex_api import (
    analyze_column,
    detect_column_type,
    query_description_generation,
    query_type_classification,
    make_json_serializable
)

# Import DQV export functionality
from dqv_export import create_dqv_metadata

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 20 * 1024 * 1024  # 16MB max file size
app.config['UPLOAD_FOLDER'] = tempfile.gettempdir()

# Store session data in memory (in production, use Redis or database)
sessions = {}


def extract_tables_from_docx(path):
    """Extract tables from DOCX files"""
    try:
        from docx import Document
        doc = Document(path)
        table_texts = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    rows.append(" | ".join(cells))
            table_texts.append("\n".join(rows))
        return "\n\n".join(table_texts)
    except Exception as e:
        return f"[Error extracting tables: {e}]"


def read_extra_file(file_path):
    """Read and process additional context files"""
    try:
        if file_path.endswith(".txt"):
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        elif file_path.endswith(".json"):
            with open(file_path, 'r', encoding='utf-8') as f:
                return json.dumps(json.load(f), indent=2)
        elif file_path.endswith(".pdf"):
            try:
                import fitz  # PyMuPDF
                doc = fitz.open(file_path)
                text_parts = []
                for page in doc:
                    text_parts.append(page.get_text())
                doc.close()
                return "\n".join(text_parts)
            except ImportError:
                return "[PDF processing requires PyMuPDF library]"
        elif file_path.endswith(".docx"):
            try:
                from docx import Document
                doc = Document(file_path)
                text_parts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
                table_text = extract_tables_from_docx(file_path)
                return "\n\n".join(text_parts + [table_text])
            except ImportError:
                return "[DOCX processing requires python-docx library]"
        elif file_path.endswith(".xlsx"):
            try:
                table_text = extract_tables_from_docx(file_path)
                return "\n\n".join(table_text)
                #df = pd.read_excel(file_path)
                #return df.to_csv(index=False)
            except Exception as e:
                return f"[Error reading Excel file: {e}]"
        elif file_path.endswith(".csv"):
            try:
                table_text = extract_tables_from_docx(file_path)
                return "\n\n".join(table_text)
                #df = pd.read_excel(file_path)
                #return df.to_csv(index=False)
            except Exception as e:
                return f"[Error reading Excel file: {e}]"



        else:
            return f"[Unsupported file format: {file_path}]"
    except Exception as e:
        return f"[Error reading extra file: {e}]"


def get_prompt_context(extra_content):
    """Format additional context for prompts"""
    if extra_content and extra_content.strip():
        return f"\n\nAdditional file context:\n{extra_content}\n"
    return ""


@app.route('/')
def index():
    """Serve the main web interface"""
    return render_template('index.html')


@app.route('/upload', methods=['POST'])
def upload_file():
    """Handle CSV file upload and optional additional file"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No CSV file provided'}), 400

        csv_file = request.files['file']
        if csv_file.filename == '':
            return jsonify({'error': 'No CSV file selected'}), 400

        if not csv_file.filename.lower().endswith('.csv'):
            return jsonify({'error': 'Main file must be a CSV'}), 400

        # Read CSV data
        try:
            dataset = pd.read_csv(csv_file)
        except Exception as e:
            return jsonify({'error': f'Error reading CSV file: {str(e)}'}), 400

        if dataset.empty:
            return jsonify({'error': 'CSV file is empty'}), 400

        # Generate session ID
        session_id = datetime.now().strftime('%Y%m%d_%H%M%S_%f')

        # Initialize session data
        session_data = {
            'dataset': dataset,
            'filename': secure_filename(csv_file.filename),
            'columns_processed': [],
            'metadata': {},
            'extra_content': '',
            'analysis_results': {}  # Track analysis results for previous columns context
        }

        # Handle optional additional file
        if 'extra_file' in request.files:
            extra_file = request.files['extra_file']
            if extra_file.filename != '':
                # Save the extra file temporarily
                extra_filename = secure_filename(extra_file.filename)
                extra_filepath = os.path.join(app.config['UPLOAD_FOLDER'], f"{session_id}_{extra_filename}")
                extra_file.save(extra_filepath)

                # Process the extra file
                extra_content = read_extra_file(extra_filepath)
                session_data['extra_content'] = extra_content
                session_data['extra_filename'] = extra_filename

                # Clean up the temporary file
                try:
                    os.remove(extra_filepath)
                except:
                    pass

        # Store dataset in session
        sessions[session_id] = session_data

        # Return file info and preview
        preview_data = dataset.head().fillna('').to_dict('records')

        response_data = {
            'session_id': session_id,
            'filename': csv_file.filename,
            'rows': len(dataset),
            'columns': len(dataset.columns),
            'column_names': list(dataset.columns),
            'preview': preview_data
        }

        # Add extra file info if present
        if session_data['extra_content']:
            response_data['extra_file'] = {
                'filename': session_data.get('extra_filename', 'unknown'),
                'content_length': len(session_data['extra_content']),
                'content_preview': session_data['extra_content'][:200] + '...' if len(
                    session_data['extra_content']) > 200 else session_data['extra_content']
            }

        return jsonify(response_data)

    except Exception as e:
        return jsonify({'error': f'Error processing files: {str(e)}'}), 500


@app.route('/set_dataset_info', methods=['POST'])
def set_dataset_info():
    """Store dataset name and description"""
    try:
        data = request.json
        session_id = data.get('session_id')

        if session_id not in sessions:
            return jsonify({'error': 'Invalid session'}), 400

        sessions[session_id]['metadata'].update({
            'dataset_name': data.get('name'),
            'dataset_description': data.get('description')
        })

        return jsonify({'success': True})

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/analyze_column', methods=['POST'])
def analyze_column_endpoint():
    """Analyze a specific column with AI assistance - Two separate LLM calls"""
    try:
        data = request.json
        session_id = data.get('session_id')
        column_name = data.get('column_name')

        if session_id not in sessions:
            return jsonify({'error': 'Invalid session'}), 400

        session_data = sessions[session_id]
        dataset = session_data['dataset']

        if column_name not in dataset.columns:
            return jsonify({'error': 'Column not found'}), 400

        # Analyze column using existing functions
        column_data = dataset[column_name]

        # Ensure we have valid data
        if column_data.empty:
            return jsonify({'error': f'Column {column_name} is empty'}), 400

        try:
            stats = analyze_column(column_data)
            detected_type = detect_column_type(column_data)
        except Exception as e:
            return jsonify({'error': f'Error analyzing column data: {str(e)}'}), 500

        # Get dataset context for LLM
        dataset_name = session_data['metadata'].get('dataset_name', 'Unknown Dataset')
        dataset_description = session_data['metadata'].get('dataset_description', 'No description')
        dataset_sample_str = dataset.head().to_string(index=False)

        # Get previously analyzed columns from session (temporary analysis results)
        previous_columns = []
        if 'analysis_results' in session_data:
            for col_name, col_data in session_data['analysis_results'].items():
                if col_name != column_name and col_data.get('description') and col_data.get('type'):
                    previous_columns.append({
                        'name': col_name,
                        'type': col_data.get('type'),
                        'description': col_data.get('description')
                    })

        # Get additional context from extra file
        extra_content = session_data.get('extra_content', '')
        prompt_context = get_prompt_context(extra_content)

        print(f"\n🔄 Starting analysis for column: {column_name}")
        if extra_content:
            print(f"📄 Using additional context from: {session_data.get('extra_filename', 'unknown file')}")

        # FIRST LLM CALL: Generate description
        print("📝 LLM Call #1: Generating description...")
        try:
            description = query_description_generation(
                column_name, stats, detected_type,
                dataset_name, dataset_description, dataset_sample_str,
                previous_columns, prompt_context
            )
            print(f"✅ Description generated: {description[:100]}...")
        except Exception as e:
            print(f"❌ LLM description failed: {e}")
            description = f"This column represents {column_name} data in the dataset."

        # SECOND LLM CALL: Type classification using the description
        print("🎯 LLM Call #2: Classifying type based on description...")
        try:
            llm_result = query_type_classification(
                column_name, description, stats, detected_type,
                dataset_name, dataset_description, dataset_sample_str,
                prompt_context
            )
            type_confidence = llm_result.get('confidence', {})
            suggested_type = llm_result.get('suggested_type', detected_type)
            print(f"✅ Type classified as: {suggested_type}")
        except Exception as e:
            print(f"❌ LLM type classification failed: {e}")
            # Fallback to rule-based detection
            fallback_confidence = {
                "binary": 0.0, "categorical": 0.0, "ordinal": 0.0,
                "continuous": 0.0, "identifier": 0.0, "free_text": 0.0
            }
            fallback_confidence[detected_type] = 0.9
            type_confidence = fallback_confidence
            suggested_type = detected_type

        # Clean stats for JSON serialization
        clean_stats = json.loads(json.dumps(stats, default=make_json_serializable))

        # Store analysis results in session for future reference
        if 'analysis_results' not in session_data:
            session_data['analysis_results'] = {}

        session_data['analysis_results'][column_name] = {
            'description': description,
            'type': suggested_type,
            'stats': clean_stats
        }

        return jsonify({
            'column_name': column_name,
            'stats': clean_stats,
            'detected_type': detected_type,
            'suggested_description': description,
            'suggested_type': suggested_type,
            'type_confidence': type_confidence
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/reanalyze_type', methods=['POST'])
def reanalyze_type():
    """Reanalyze column type based on updated description - Third LLM call"""
    try:
        data = request.json
        session_id = data.get('session_id')
        column_name = data.get('column_name')
        updated_description = data.get('description')

        if session_id not in sessions:
            return jsonify({'error': 'Invalid session'}), 400

        session_data = sessions[session_id]
        dataset = session_data['dataset']

        if column_name not in dataset.columns:
            return jsonify({'error': 'Column not found'}), 400

        # Get column stats (reuse existing analysis)
        column_data = dataset[column_name]
        stats = analyze_column(column_data)
        detected_type = detect_column_type(column_data)

        # Get dataset context
        dataset_name = session_data['metadata'].get('dataset_name', 'Unknown Dataset')
        dataset_description = session_data['metadata'].get('dataset_description', 'No description')
        dataset_sample_str = dataset.head().to_string(index=False)

        # Get previously analyzed columns from session
        previous_columns = []
        if 'analysis_results' in session_data:
            for col_name, col_data in session_data['analysis_results'].items():
                if col_name != column_name and col_data.get('description') and col_data.get('type'):
                    previous_columns.append({
                        'name': col_name,
                        'type': col_data.get('type'),
                        'description': col_data.get('description')
                    })

        # Get additional context from extra file
        extra_content = session_data.get('extra_content', '')
        prompt_context = get_prompt_context(extra_content)

        print(f"\n🔄 Reanalyzing type for column: {column_name}")
        print(f"📝 Updated description: {updated_description[:100]}...")

        # THIRD LLM CALL: Query LLM for type classification using the updated description
        print("🎯 LLM Call #3: Reclassifying type based on updated description...")
        try:
            llm_result = query_type_classification(
                column_name, updated_description, stats, detected_type,
                dataset_name, dataset_description, dataset_sample_str,
                prompt_context
            )
            suggested_type = llm_result.get('suggested_type', detected_type)
            print(f"✅ Type reclassified as: {suggested_type}")

            # Update stored analysis results
            if 'analysis_results' not in session_data:
                session_data['analysis_results'] = {}
            if column_name in session_data['analysis_results']:
                session_data['analysis_results'][column_name]['description'] = updated_description
                session_data['analysis_results'][column_name]['type'] = suggested_type
        except Exception as e:
            print(f"❌ LLM type classification failed: {e}")
            # Fallback to rule-based detection
            suggested_type = detected_type

        return jsonify({
            'column_name': column_name,
            'suggested_type': suggested_type
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/confirm_column', methods=['POST'])
def confirm_column():
    """Save user-confirmed column metadata"""
    try:
        data = request.json
        session_id = data.get('session_id')
        column_name = data.get('column_name')
        column_type = data.get('type')
        description = data.get('description')

        if session_id not in sessions:
            return jsonify({'error': 'Invalid session'}), 400

        session_data = sessions[session_id]
        dataset = session_data['dataset']

        # Get column stats
        column_data = dataset[column_name]
        stats = analyze_column(column_data)

        # Create column metadata entry
        column_entry = {
            "name": column_name,
            "description": description,
            "type": column_type,
            "missing_values": int(stats.get("missing_values", 0)),
            "unique_values": int(stats.get("unique_values", 0))
        }

        # Add numerical stats for continuous columns
        if column_type == "continuous" and stats.get("mean") is not None:
            column_entry.update({
                "mean": float(stats.get("mean")),
                "std": float(stats.get("std")) if stats.get("std") is not None else None,
                "min": float(stats.get("min")) if stats.get("min") is not None else None,
                "max": float(stats.get("max")) if stats.get("max") is not None else None
            })

        # Add to processed columns
        session_data['columns_processed'].append(column_entry)

        return jsonify({'success': True})

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/get_metadata', methods=['POST'])
def get_metadata():
    """Retrieve complete metadata for the dataset"""
    try:
        data = request.json
        session_id = data.get('session_id')

        if session_id not in sessions:
            return jsonify({'error': 'Invalid session'}), 400

        session_data = sessions[session_id]

        metadata = {
            "dataset_name": session_data['metadata'].get('dataset_name'),
            "dataset_description": session_data['metadata'].get('dataset_description'),
            "columns": session_data['columns_processed']
        }

        return jsonify(metadata)

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/download_metadata', methods=['POST'])
def download_metadata():
    """Generate and download metadata JSON file"""
    try:
        data = request.json
        session_id = data.get('session_id')
        format_type = data.get('format', 'json')  # 'json' or 'dqv'

        if session_id not in sessions:
            return jsonify({'error': 'Invalid session'}), 400

        session_data = sessions[session_id]

        metadata = {
            "dataset_name": session_data['metadata'].get('dataset_name'),
            "dataset_description": session_data['metadata'].get('dataset_description'),
            "columns": session_data['columns_processed']
        }

        # Clean metadata for JSON serialization
        clean_metadata = json.loads(json.dumps(metadata, default=make_json_serializable))

        # Get dataset name and create safe filename
        dataset_name = metadata.get('dataset_name', 'dataset')
        safe_dataset_name = dataset_name.replace(' ', '_').replace('-', '_').lower()
        # Remove any other potentially problematic characters
        import re
        safe_dataset_name = re.sub(r'[^\w\-_]', '', safe_dataset_name)

        if format_type == 'dqv':
            # Generate DQV format
            try:
                dqv_content = create_dqv_metadata(clean_metadata)
                filename = f"{safe_dataset_name}_metadata.ttl"
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], f"{session_id}_{filename}")

                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(dqv_content)

                return send_file(
                    filepath,
                    as_attachment=True,
                    download_name=filename,
                    mimetype='text/turtle'
                )

            except Exception as e:
                return jsonify({'error': f'Error generating DQV format: {str(e)}'}), 500
        else:
            # Generate JSON format (default)
            filename = f"{safe_dataset_name}_metadata.json"
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], f"{session_id}_{filename}")

            with open(filepath, 'w') as f:
                json.dump(clean_metadata, f, indent=4)

            return send_file(
                filepath,
                as_attachment=True,
                download_name=filename,
                mimetype='application/json'
            )

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({
        'status': 'healthy',
        'active_sessions': len(sessions),
        'version': '1.1.3',
        'supported_formats': ['json', 'dqv'],
        'supported_extra_files': ['.txt', '.json', '.pdf', '.docx']
    })


@app.errorhandler(413)
def too_large(e):
    """Handle file too large error"""
    return jsonify({'error': 'File too large. Maximum size is 20MB.'}), 413


@app.errorhandler(404)
def not_found(e):
    """Handle 404 errors"""
    return jsonify({'error': 'Endpoint not found'}), 404


@app.errorhandler(500)
def internal_error(e):
    """Handle internal server errors"""
    return jsonify({'error': 'Internal server error'}), 500


# Cleanup function to remove old sessions (optional)
def cleanup_old_sessions():
    """Remove sessions older than 1 hour to prevent memory leaks"""
    import time
    current_time = time.time()
    cutoff_time = current_time - 3600  # 1 hour ago

    sessions_to_remove = []
    for session_id in sessions:
        try:
            # Extract timestamp from session_id
            timestamp_str = session_id.split('_')[0] + session_id.split('_')[1]
            session_time = datetime.strptime(timestamp_str, '%Y%m%d%H%M%S').timestamp()
            if session_time < cutoff_time:
                sessions_to_remove.append(session_id)
        except:
            # If we can't parse the timestamp, remove the session
            sessions_to_remove.append(session_id)

    for session_id in sessions_to_remove:
        del sessions[session_id]
        print(f"Cleaned up old session: {session_id}")


if __name__ == '__main__':
    print("=" * 60)
    print("Dataset Metadata Extraction Tool - Web Interface")
    print("=" * 60)
    print("Features:")
    print("   • AI-powered column description generation (LLM Call #1)")
    print("   • AI-powered column type detection (LLM Call #2)")
    print("   • Interactive description updates (LLM Call #3)")
    print("   • Full additional file context support (no truncation)")
    print("   • Column navigation (back/forth between columns)")
    print("   • Professional metadata export (JSON & DQV formats)")
    print()
    print("🔧 Configuration:")
    print(f"   • Max file size: {app.config['MAX_CONTENT_LENGTH'] // 1024 // 1024}MB")
    print("   • LLM API: Configured in meta_data_ex_api.py")
    print("   • Export formats: JSON, DQV (W3C Data Quality Vocabulary)")
    print("   • Supported extra files: .txt, .json, .pdf, .docx")
    print()
    print("Starting server...")
    print("   • Web interface: http://localhost:5000")
    print("   • Health check: http://localhost:5000/health")
    print()
    print("⚠️  Make sure your LLM server is running for AI features!")
    print("📄 Optional: Upload additional context files (.txt, .json, .pdf, .docx)")
    print("💡 No content truncation - full file context will be used")
    print("=" * 60)

    # Run the Flask application
    app.run(debug=True, host='0.0.0.0', port=5000)
